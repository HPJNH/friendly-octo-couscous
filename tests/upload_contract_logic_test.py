from pathlib import Path
import sys
from tempfile import TemporaryDirectory

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from app import create_app
from app.constants import (
    DRAFT_MAIN_TITLE,
    DRAFT_STANDARD_SUBSECTIONS,
    DRAFT_STRUCTURED_TABLE_COLUMNS,
    DRAFT_STRUCTURED_TABLE_COLUMNS_V2,
    SECTION_DEFINITIONS,
)
from app.parsers import parse_draft_file, validate_draft_contract
from app.rebuild_engine import (
    entry_to_card,
    extract_entries_from_payload,
    find_linked_matches,
    has_meaningful_delta,
)
from app.rendering import build_section_render_payload, extract_cards_from_render_payload
from app.services import extract_card_business_tags
from app.utils import normalize_compare_text


OLD_SUBSECTION_TITLES = {
    "1": "补采窗口内新增信号",
    "2": "近72小时重点新信号",
    "3": "背景补充",
    "4": "当前状态说明",
}
NEW_SUBSECTION_TITLES = dict(DRAFT_STANDARD_SUBSECTIONS)


def add_table(document: Document, headers: list[str], row_values: list[str]) -> None:
    table = document.add_table(rows=2, cols=len(headers))
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
    for index, value in enumerate(row_values):
        table.cell(1, index).text = value


def build_full_draft(
    path: Path,
    *,
    headers: list[str],
    subsection_titles: dict[str, str],
    include_explicit_metadata: bool,
    real_row: list[str],
) -> None:
    document = Document()
    document.add_paragraph(DRAFT_MAIN_TITLE)

    for definition in SECTION_DEFINITIONS:
        number = definition["number"]
        title = definition["title"]
        document.add_paragraph(f"{number}. {title}")

        if definition["key"] == "report_note":
            document.add_paragraph("1.1 数据与规则")
            if include_explicit_metadata:
                document.add_paragraph("本期新增窗口：2026-04-10 至 2026-04-12")
                document.add_paragraph("近72小时窗口：2026-04-10 至 2026-04-12")
            else:
                document.add_paragraph("本期补采以近三日行业变化为主，按既有阅读规则继续处理。")
            document.add_paragraph("1.2 编制原则")
            document.add_paragraph("历史参考：2025-03-01 的旧规则仅用于说明，不应拖偏当前窗口。")
            continue

        for suffix in ("1", "2", "3", "4"):
            document.add_paragraph(f"{number}.{suffix} {subsection_titles[suffix]}")
            if suffix == "4":
                document.add_paragraph(f"{title} 当前以结构化跟踪为主，本期说明仅用于辅助阅读。")
                continue

            if definition["key"] == "themed_tracks" and suffix == "1":
                add_table(document, headers, real_row)
                continue

            add_table(document, headers, build_placeholder_row(headers, suffix))

    document.save(path)


def build_placeholder_row(headers: list[str], suffix: str) -> list[str]:
    if len(headers) == len(DRAFT_STRUCTURED_TABLE_COLUMNS):
        title = {
            "1": "本期未监测到有效新增信号",
            "2": "本重点窗口内未监测到额外重点新信号",
            "3": "本模块未启用背景补充",
        }[suffix]
        return [title, "", "", "", "", ""]

    title = {
        "1": "本期未监测到有效新增信号",
        "2": "本重点窗口内未监测到额外重点新信号",
        "3": "本模块未启用背景补充",
    }[suffix]
    return [title, "", "", "", "", "", "", "", ""]


def build_fake_row(path: Path, *, report_date: str = "2026-04-12") -> dict:
    return {
        "id": 1,
        "original_name": path.name,
        "stored_path": str(path),
        "report_date": report_date,
    }


def make_entry(
    contract_version: str,
    *,
    source_level: str = "A2",
    time_text: str = "2026-04-11",
    delta_text: str = "",
    core_content: str = "发布空间香新品套装",
) -> dict:
    return {
        "title": "空间香新品套装",
        "core_content": core_content,
        "why_included": "可作为主题赛道样本",
        "source_level": source_level,
        "time_text": time_text,
        "event_date": "2026-04-11",
        "source_name": "品牌官网",
        "evidence_json": {
            "contract_version": contract_version,
            "delta_text": delta_text,
            "business_tags": ["space_home_fragrance"],
        },
    }


def assert_v2_contract_flow(path: Path) -> None:
    validation = validate_draft_contract(path)
    assert validation["success"], validation["errors"]

    payload = parse_draft_file(path)
    assert payload["metadata"]["contract_version"] == "v2"
    assert payload["metadata"]["window_metadata"]["patch_window"] == ["2026-04-10", "2026-04-12"]
    assert payload["metadata"]["window_metadata"]["focus_window"] == ["2026-04-10", "2026-04-12"]

    section_payload = payload["sections"]["themed_tracks"]
    render_payload = build_section_render_payload(
        "themed_tracks",
        "主题型赛道观察",
        section_payload["blocks"],
        "历史保留",
    )
    cards = extract_cards_from_render_payload(render_payload)
    real_card = next(card for card in cards if normalize_compare_text(card.get("title", "")) == normalize_compare_text("空间香新品套装"))
    assert real_card["source_url"] == "https://example.com/theme-space"
    assert real_card["delta_text"] == "新增发布时间、上线节奏与套装规格。"
    assert real_card["business_tags"] == ["space_home_fragrance"]
    assert extract_card_business_tags(real_card) == ["space_home_fragrance"]

    extracted_entries, extraction_notes = extract_entries_from_payload(build_fake_row(path), payload)
    assert extraction_notes == [], extraction_notes
    themed_entry = next(
        entry
        for entry in extracted_entries
        if entry["module_key"] == "themed_tracks" and entry["entry_type"] == "real"
    )
    assert themed_entry["source_url"] == "https://example.com/theme-space"
    assert themed_entry["evidence_json"]["delta_text"] == "新增发布时间、上线节奏与套装规格。"
    assert themed_entry["evidence_json"]["business_tags"] == ["space_home_fragrance"]
    assert themed_entry["is_in_patch_window"] == 1
    assert themed_entry["is_in_focus_window"] == 1

    card_from_entry = entry_to_card(
        {
            "id": 99,
            "title": themed_entry["title"],
            "display_status": themed_entry["display_status"],
            "time_text": themed_entry["time_text"],
            "source_name": themed_entry["source_name"],
            "core_content": themed_entry["core_content"],
            "why_included": themed_entry["why_included"],
            "source_level": themed_entry["source_level"],
            "section_type": themed_entry["section_type"],
            "entry_type": themed_entry["entry_type"],
            "source_url": themed_entry["source_url"],
            "source_title": themed_entry["source_title"],
            "supporting_sources_json": themed_entry["supporting_sources_json"],
            "first_seen_date": themed_entry["first_seen_date"],
            "last_seen_date": themed_entry["last_seen_date"],
            "confidence_level": themed_entry["confidence_level"],
            "needs_review": themed_entry["needs_review"],
            "event_key": themed_entry["event_key"],
            "report_date": themed_entry["report_date"],
            "evidence_json": themed_entry["evidence_json"],
        }
    )
    assert card_from_entry["business_tags"] == ["space_home_fragrance"]
    assert card_from_entry["compare_meta"]["business_tags"] == ["space_home_fragrance"]
    assert card_from_entry["delta_text"] == "新增发布时间、上线节奏与套装规格。"


def assert_v1_contract_flow(path: Path) -> None:
    validation = validate_draft_contract(path)
    assert validation["success"], validation["errors"]

    payload = parse_draft_file(path)
    assert payload["metadata"]["contract_version"] == "v1"
    assert payload["metadata"]["window_metadata"] == {}

    extracted_entries, extraction_notes = extract_entries_from_payload(build_fake_row(path), payload)
    assert extraction_notes == [], extraction_notes
    themed_entry = next(
        entry
        for entry in extracted_entries
        if entry["module_key"] == "themed_tracks" and entry["entry_type"] == "real"
    )
    assert themed_entry["source_url"] == ""
    assert themed_entry["evidence_json"]["delta_text"] == ""
    assert themed_entry["evidence_json"]["business_tags"] == []


def main() -> None:
    with TemporaryDirectory() as temp_dir:
        temp_root = Path(temp_dir)
        v2_path = temp_root / "sample_v2_new_headings.docx"
        build_full_draft(
            v2_path,
            headers=DRAFT_STRUCTURED_TABLE_COLUMNS_V2,
            subsection_titles=NEW_SUBSECTION_TITLES,
            include_explicit_metadata=True,
            real_row=[
                "空间香新品套装",
                "2026-04-11",
                "A2",
                "品牌官网",
                "https://example.com/theme-space",
                "发布空间香新品套装，并公布上线节奏。",
                "新增发布时间、上线节奏与套装规格。",
                "空间香 / 家居香氛",
                "有助于识别主题赛道的新产品表达方向。",
            ],
        )
        assert_v2_contract_flow(v2_path)

        v1_path = temp_root / "sample_v1_old_headings.docx"
        build_full_draft(
            v1_path,
            headers=DRAFT_STRUCTURED_TABLE_COLUMNS,
            subsection_titles=OLD_SUBSECTION_TITLES,
            include_explicit_metadata=False,
            real_row=[
                "观夏空间香新品",
                "2026-04-11",
                "A2",
                "品牌官网",
                "发布空间香新品套装，并公布上线节奏。",
                "作为主题赛道中的参考样本，值得保留。",
            ],
        )
        assert_v1_contract_flow(v1_path)

    v2_previous = make_entry("v2", delta_text="")
    v2_same = make_entry("v2", delta_text="")
    assert not has_meaningful_delta(v2_previous, v2_same)

    v2_with_delta = make_entry("v2", delta_text="新增联名信息")
    assert has_meaningful_delta(v2_previous, v2_with_delta)

    v2_source_gain = make_entry("v2", source_level="A1", delta_text="")
    assert has_meaningful_delta(v2_previous, v2_source_gain)

    v1_previous = make_entry("v1", delta_text="", core_content="发布空间香新品套装")
    v1_changed = make_entry(
        "v1",
        delta_text="",
        core_content="海外渠道反馈出现完全不同的联名陈列事件，并带来新的成交与复购信号。",
    )
    v1_changed["why_included"] = "已从新品发布转成渠道成交验证，属于旧版相似度回退场景。"
    assert has_meaningful_delta(v1_previous, v1_changed)

    app = create_app()
    with app.app_context():
        exact_bucket = normalize_compare_text("空间香新品套装")
        near_bucket = normalize_compare_text("空间香新品套装（扩展标题）")
        grouped_registry = {
            exact_bucket: [
                {
                    "title": "空间香新品套装",
                    "source_name": "品牌官网",
                    "source_url": "https://example.com/exact",
                    "time_text": "2026-04-11",
                    "core_content": "官方发布新品套装信息。",
                    "source_level": "A2",
                },
                {
                    "title": "空间香新品套装",
                    "source_name": "品牌官网",
                    "source_url": "https://example.com/older",
                    "time_text": "2026-04-08",
                    "core_content": "更早版本说明。",
                    "source_level": "A2",
                },
            ],
            near_bucket: [
                {
                    "title": "空间香新品套装（扩展标题）",
                    "source_name": "聚合站",
                    "source_url": "https://example.com/fuzzy",
                    "time_text": "2026-04-11",
                    "core_content": "模糊标题命中的聚合内容。",
                    "source_level": "C",
                }
            ],
        }

        item = {
            "title": "空间香新品套装",
            "core_content": "发布空间香新品套装",
            "event_date": "2026-04-11",
            "source_name": "品牌官网",
        }
        matches = find_linked_matches(item, grouped_registry)
        assert len(matches) == 1
        assert matches[0]["source_url"] == "https://example.com/exact"

        fuzzy_only_item = {
            "title": "空间香新品",
            "core_content": "发布空间香新品套装",
            "event_date": "2026-04-11",
            "source_name": "品牌官网",
        }
        assert find_linked_matches(fuzzy_only_item, grouped_registry) == []

    print("upload_contract_logic_test_ok")


if __name__ == "__main__":
    main()
